# That's a script, okey? That's why we use some fucking shitty decisions.
import datetime
import os
import sys
import docx
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from openpyxl import load_workbook
from matplotlib import pyplot


# https://stackoverflow.com/questions/55115070/next-letter-in-alphabet
def nextAlphabetLetter(s):
    return chr((ord(s.upper()) + 1 - 65) % 26 + 65)


def throwWrongStructure():
    raise Exception("Конфигурационный файл не соответствует требуемому формату")


# fTest: python test.py "C:\Users\dobri\Desktop\Свод проектов 529.xlsx" . C:\Users\dobri\Desktop\ReportAutomation\app\src\config\stdConfig.ycfg 1
def goToNextSharps(lines, curLineNumber):
    while curLineNumber < len(lines):
        if lines[curLineNumber][:3] == "###":
            return curLineNumber
        curLineNumber += 1
    return -1


'''
    Function forms rules in a dictionary for the next step (excel table processing)
    
    Forms a dictionary with fields: startOfTable, headers.
    --- startOfTable contains char and int which are coordinates of start of table (horizontal and then vertical)
    --- headers contains two strings: format of field, it's horizontal letter (coordinate in excel lol ok)
'''


def getRulesFromConfig(configPath):
    outRules = dict()

    with open(configPath, encoding = 'utf-8') as file:
        lines = file.readlines()

    # Search of table start coordinates
    curLineNumber = goToNextSharps(lines, 0)
    # print(lines[curLineNumber].replace(" \n\0", ""))
    if lines[curLineNumber] != "### Область названий организаций\n":
        throwWrongStructure()

    curLineNumber += 1
    outRules["startOfTable"] = lines[curLineNumber].replace("\n", "").split(" ")
    outRules["startOfTable"][1] = int(outRules["startOfTable"][1])

    #    for i in outRules["startOfTable"]:
    #        print("!!{}".format(i))

    # Search of headers in configFile
    curLineNumber = goToNextSharps(lines, curLineNumber)
    curLineNumber += 1
    if curLineNumber == len(lines):
        throwWrongStructure()

    headersInDocument = []
    curHeaderDescription = lines[curLineNumber]
    while curLineNumber < len(lines) and curHeaderDescription[0][0] == "[":
        curHeaderDescription = lines[curLineNumber].replace("\n", "").split(" ")
        # print(curHeaderDescription)
        headersInDocument.append(curHeaderDescription)
        curLineNumber += 1

    outRules["headers"] = headersInDocument

    return outRules


'''
    Function returns information from table which it gets using rules
    
    Rules specification is in the function getRulesFromConfig
'''


def getInfoFromExcelTableUsingRules(excelTablePath, rules, rowNumber):
    gotData = rules["headers"]
    # print(gotData)
    workbook = load_workbook(filename = excelTablePath, data_only = True)
    sheets = workbook.sheetnames
    ranges = workbook[sheets[0]]

    headersStartIndexNumber = rules["startOfTable"][1] - 1
    processingRowIndexNumber = rowNumber + rules["startOfTable"][1] - 1
    # print(ranges[rules["startOfTable"][0] + str(processingRowIndexNumber)].value)
    # print(ranges[gotData[0][1] + str(processingRowIndexNumber)].value)
    # print(ranges[gotData[0][1] + str(headersStartIndexNumber)].value)

    i = 0
    diagramCounter = 0
    while i < len(gotData):
        if gotData[i][0] == "[текст]":
            gotData[i][0] = ranges[gotData[i][1] + str(headersStartIndexNumber)].value

            gotInfo = ranges[gotData[i][1] + str(processingRowIndexNumber)].value
            if type(gotInfo) == datetime.datetime:
                gotInfo = "{}.{}.{}".format(gotInfo.day, gotInfo.month, gotInfo.year)

            gotData[i][1] = gotInfo
        elif gotData[i][0] == "[столбчатаядиаграмма]":
            gotData[i][0] = "Диаграмма по выплатам {}".format(diagramCounter)
            diagramCounter += 1

            # Damn.
            namesOfColumnsInDiagram = gotData[i][1].split("-")
            endLetter = namesOfColumnsInDiagram[1]
            namesOfColumnsInDiagram = [namesOfColumnsInDiagram[0]]
            curLetter = namesOfColumnsInDiagram[0]
            while curLetter < endLetter:
                curLetter = nextAlphabetLetter(curLetter)
                namesOfColumnsInDiagram.append(curLetter)

            groups = []
            counts = []
            for curColumn in namesOfColumnsInDiagram:
                groups.append(ranges[curColumn + str(headersStartIndexNumber)].value)
                counts.append(ranges[curColumn + str(processingRowIndexNumber)].value)

            # TODO To separated function
            pyplot.bar(groups, counts,
                       color = ['paleturquoise', 'aquamarine', 'mediumspringgreen', 'mediumseagreen', 'green',
                                'darkgreen'])
            pyplot.ticklabel_format(axis = "y", scilimits = (6, 6))
            pyplot.xlabel("Период")
            pyplot.ylabel("Выплаты, млн. руб.")
            pyplot.yticks(counts)
            pyplot.grid(axis = 'y', linestyle = "-")

            if not os.path.exists("ImagesForYcfg"):
                os.mkdir("ImagesForYcfg")
            imgPath = "ImagesForYcfg/figPyYcfg{}.png".format(diagramCounter)
            pyplot.savefig(imgPath, bbox_inches = 'tight', dpi = 100)
            gotData[i][1] = imgPath
        i += 1

    gotData.insert(0, ranges[rules["startOfTable"][0] + str(processingRowIndexNumber)].value)

    # print("After")
    # print(gotData)

    return gotData


def setParaFormatPlainText(format, font):
    format.first_line_indent = Cm(1.25)
    format.line_spacing = 1
    format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    font.name = "Times New Roman"
    font.size = Pt(14)


def setParaFormatHeading(format, font):
    setParaFormatPlainText(format, font)
    font.bold = True


def setParaFormatTitle(format, font):
    format.line_spacing = 1
    format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    font.name = "Times New Roman"
    font.bold = True
    font.size = Pt(16)


def setPictureFormat(format, run, picPath):
    format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run.add_picture(picPath, width = Cm(10))


def formDocxFile(gotData, savePath):
    doc = docx.Document()
    filename = "Отчёт по {}".format(gotData[0])
    for s in ['*', '.', '/', '\\', '[', ']', ':', ';', '|', ',']:
        filename = filename.replace(s, '')
    filename += ".docx"
    filename = filename.replace('"', "'")
    filename = savePath + "/" + filename

    titlePara = doc.add_paragraph()
    setParaFormatTitle(titlePara.paragraph_format,
                       titlePara.add_run("Аналитический отчёт по комплексному проекту {}".format(gotData[0])).font)
    gotData.pop(0)
    print(gotData)
    for currentDataPair in gotData:
        if currentDataPair[0][:9] == "Диаграмма":
            headerPara = doc.add_paragraph()
            setParaFormatHeading(headerPara.paragraph_format, headerPara.add_run("{}".format(currentDataPair[0])).font)

            picturePara = doc.add_paragraph()
            setPictureFormat(picturePara.paragraph_format, picturePara.add_run(), currentDataPair[1])
        else:
            headerPara = doc.add_paragraph()
            setParaFormatHeading(headerPara.paragraph_format, headerPara.add_run("{}".format(currentDataPair[0])).font)

            plainTextPara = doc.add_paragraph()
            setParaFormatPlainText(plainTextPara.paragraph_format,
                                   plainTextPara.add_run("{}".format(currentDataPair[1].replace("\n", " "))).font)

    doc.save(filename)


'''
    REALLY FUCKING IMPORTANT!!!
    We have a small trouble using some names with spaces. That's why
    we MUST send ways in quotes like 'C:/Go Fuck Yourself/Program Files/fuck fuck fuck fuck'
    Why can't we use '_'? BECAUSE OF FUCKING 'FUCK:/GO_FUCK/FUCKING_FUCK/FUCK FUCK FUCK'
    
    I think you've got it, right?
'''


def main():
    gotArgs = sys.argv

    # try:
    tablePath = gotArgs[3]  # path to table to get info from
    docSavePath = gotArgs[4]  # path to save document
    configPath = gotArgs[5]  # path to configuration
    workRowNumber = int(gotArgs[6])  # number of row from table
    # except:
    #     return 1

    dictOfRulesForDoc = getRulesFromConfig(configPath)
    gotInfo = getInfoFromExcelTableUsingRules(tablePath, dictOfRulesForDoc, workRowNumber)
    formDocxFile(gotInfo, docSavePath)

    return 0


ret = main()
if not ret:
    print("OK")
else:
    print("err")
