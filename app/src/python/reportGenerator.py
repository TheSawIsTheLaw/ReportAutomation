# That's a script, okey? That's why we use some fucking shitty decisions.
import datetime
import os
import sys
import docx
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import locale

from openpyxl.utils import get_column_letter
from openpyxl.utils.cell import column_index_from_string
from openpyxl import load_workbook
from matplotlib import pyplot, ticker
from numpy import arange

from RGConstants import *
from floatingImage import add_float_picture

'''
    Just a raise :)
'''


def throwWrongStructure():
    raise Exception("Конфигурационный файл не соответствует требуемому формату")


'''
    Skip comments in configuration file
'''


def goToNextSharps(lines, curLineNumber):
    while curLineNumber < len(lines):
        if lines[curLineNumber][:3] == "###":
            return curLineNumber
        curLineNumber += 1
    return -1


'''
    Function forms rules in a dictionary for the next step (excel table processing)
    
    Forms a dictionary with fields: startOfTable, headers.
    --- startOfTable contains char and int which are coordinates of start of table (horizontal and then vertical)
    --- headers contains two strings: format of field, it's horizontal letter (coordinate in excel lol ok)
'''


def getRulesFromConfig(configPath):
    outRules = dict()

    with open(configPath, encoding = 'utf-8') as file:
        lines = file.readlines()

    # Search of table start coordinates
    curLineNumber = goToNextSharps(lines, 0)
    # print(lines[curLineNumber].replace(" \n\0", ""))
    if lines[curLineNumber] != "### Область названий организаций\n":
        throwWrongStructure()

    curLineNumber += 1
    outRules["startOfTable"] = lines[curLineNumber].replace("\n", "").split(" ")
    outRules["startOfTable"][1] = int(outRules["startOfTable"][1])

    #    for i in outRules["startOfTable"]:
    #        print("!!{}".format(i))

    # Search of headers in configFile
    curLineNumber = goToNextSharps(lines, curLineNumber)
    curLineNumber += 1
    if curLineNumber == len(lines):
        throwWrongStructure()

    headersInDocument = []
    curHeaderDescription = lines[curLineNumber]
    while curLineNumber < len(lines) and curHeaderDescription[0][0] == "[":
        curHeaderDescription = lines[curLineNumber].replace("\n", "").split(" ")
        # print(curHeaderDescription)
        headersInDocument.append(curHeaderDescription)
        curLineNumber += 1

    outRules["headers"] = headersInDocument
    # print(outRules)

    return outRules


'''
    Function that creates a bar plot, saves it
'''


def createBarAndSave(groups, counts, savePath):
    ax = pyplot.subplot()
    ax.bar(groups, counts,
           color = ['paleturquoise', 'aquamarine', 'mediumspringgreen', 'mediumseagreen', 'green',
                    'darkgreen'])
    ax.ticklabel_format(axis = "y", style = 'plain')
    locale.setlocale(locale.LC_ALL, '')

    def moneyAxisFormatter(x, y):
        return '{}'.format(locale.currency(x, grouping = True))

    ax.yaxis.set_major_formatter(ticker.FuncFormatter(moneyAxisFormatter))
    ax.spines.right.set_visible(False)
    ax.spines.top.set_visible(False)
    pyplot.ylabel("Величина выплат/затрат, руб.")
    pyplot.yticks(counts)
    pyplot.grid(axis = 'y', linestyle = "-")

    if not os.path.exists(PATH_TO_PYPLOT_IMAGES):
        os.mkdir(PATH_TO_PYPLOT_IMAGES)
    pyplot.savefig(savePath, bbox_inches = 'tight', dpi = 300)


'''
    Function that creates a bar plot with splited bars and saves it
'''


def createDoubleBarAndSave(groups, fcounts, scounts, savePath):
    fig, ax = pyplot.subplots(figsize = (13, 7))
    barWidth = 0.4
    xBarPos = arange(len(groups))
    fbar = ax.bar(xBarPos - barWidth / 2, fcounts, barWidth, label = 'Собственные средства', color = 'darkblue')
    sbar = ax.bar(xBarPos + barWidth / 2, scounts, barWidth, label = 'Выплаты', color = 'darkred')

    fLabels = []
    sLabels = []
    locale.setlocale(locale.LC_ALL, '')
    for i in fcounts:
        fLabels.append('{}'.format(locale.currency(i, grouping = True)))
    for i in scounts:
        sLabels.append('{}'.format(locale.currency(i, grouping = True)))

    isLong = False
    for j in fLabels:
        if len(j) >= 16:
            isLong = True
    if isLong:
        ax.bar_label(fbar, labels = fLabels, padding = 3, fontsize = 8, weight = 'bold')
    else:
        ax.bar_label(fbar, labels = fLabels, padding = 3, fontsize = 10, weight = 'bold')

    isLong = False
    for j in sLabels:
        if len(j) >= 16:
            isLong = True

    if isLong:
        ax.bar_label(sbar, labels = sLabels, padding = 3, fontsize = 8, weight = 'bold')
    else:
        ax.bar_label(sbar, labels = sLabels, padding = 3, fontsize = 10, weight = 'bold')

    pyplot.xticks(weight = 'bold', fontsize = 16)
    pyplot.yticks(weight = 'bold', fontsize = 12)

    ax.set_xticks(xBarPos)
    ax.set_xticklabels(groups)
    ax.legend(prop = {'size': 12, 'weight': 'bold'})
    ax.ticklabel_format(axis = "y", style = 'plain')

    def moneyAxisFormatter(x, y):
        return '{}'.format(locale.currency(x, grouping = True))

    ax.yaxis.set_major_formatter(ticker.FuncFormatter(moneyAxisFormatter))
    ax.spines.right.set_visible(False)
    ax.spines.top.set_visible(False)
    pyplot.ylabel("Величина выплат/затрат, руб.", weight = 'bold', fontsize = 14)

    fcounts.extend(scounts)
    fcounts.sort()
    i = 1
    # Because of fucking pyplot can not fucking space this fucking ticks
    # while i < len(fcounts):
    #     if abs(fcounts[i] - fcounts[i - 1]) < 800000:
    #         fcounts.pop(i)
    #         i -= 1
    #     i += 1

    # pyplot.yticks(fcounts, fontsize = 6)
    # pyplot.grid(axis = 'y', linestyle = "-")

    if not os.path.exists(PATH_TO_PYPLOT_IMAGES):
        os.mkdir(PATH_TO_PYPLOT_IMAGES)
    pyplot.savefig(savePath, bbox_inches = 'tight', dpi = 300)


'''
    Function returns information from table which it gets using rules
    
    Rules specification is in the function getRulesFromConfig
'''


def getInfoFromExcelTableUsingRules(excelTablePath, rules, rowNumber):
    gotData = rules["headers"]
    # print(gotData)
    workbook = load_workbook(filename = excelTablePath, data_only = True)
    sheets = workbook.sheetnames
    ranges = workbook[sheets[0]]

    headersStartIndexNumber = rules["startOfTable"][1] - 1
    processingRowIndexNumber = rowNumber + rules["startOfTable"][1]
    # print(ranges[rules["startOfTable"][0] + str(processingRowIndexNumber)].value)
    # print(ranges[gotData[0][1] + str(processingRowIndexNumber)].value)
    # print(ranges[gotData[0][1] + str(headersStartIndexNumber)].value)

    i = 0
    diagramCounter = 0
    while i < len(gotData):
        if gotData[i][0] == "[текст]":
            if "-" in gotData[i][1]:
                colRange = gotData[i][1].split("-")
                curCol = column_index_from_string(colRange[0])
                endCol = column_index_from_string(colRange[1])
                gotData.pop(i)
                while curCol <= endCol:
                    gotHeader = ranges[get_column_letter(curCol) + str(headersStartIndexNumber)].value
                    gotInfo = ranges[get_column_letter(curCol) + str(processingRowIndexNumber)].value
                    if type(gotInfo) == datetime.datetime:
                        gotInfo = "{}.{}.{}".format(gotInfo.day, gotInfo.month, gotInfo.year)
                    elif type(gotInfo) == float:  # Не забудь об этом написать в документации))
                        locale.setlocale(locale.LC_ALL, '')
                        gotInfo = '{}'.format(locale.currency(gotInfo, grouping = True))
                    gotData.insert(i, [gotHeader, gotInfo])

                    curCol += 1
                    i += 1
                i -= 1
            else:
                gotData[i][0] = ranges[gotData[i][1] + str(headersStartIndexNumber)].value

                gotInfo = ranges[gotData[i][1] + str(processingRowIndexNumber)].value
                if type(gotInfo) == datetime.datetime:
                    gotInfo = "{}.{}.{}".format(gotInfo.day, gotInfo.month, gotInfo.year)

                gotData[i][1] = gotInfo
        elif gotData[i][0] == "[столбчатаядиаграмма]":
            gotData[i][0] = "Выплаты в рамках комплексного проекта"
            diagramCounter += 1

            # Damn.
            namesOfColumnsInDiagram = gotData[i][1].split("-")
            endLetter = namesOfColumnsInDiagram[1]
            namesOfColumnsInDiagram = [namesOfColumnsInDiagram[0]]
            curLetter = namesOfColumnsInDiagram[0]
            # print(endLetter)
            curLetter = column_index_from_string(curLetter)
            # print(curLetter, column_index_from_string(endLetter))
            while curLetter < column_index_from_string(endLetter):
                curLetter += 1
                namesOfColumnsInDiagram.append(get_column_letter(curLetter))

            groups = []
            counts = []
            for curColumn in namesOfColumnsInDiagram:
                groups.append(ranges[curColumn + str(headersStartIndexNumber)].value)
                counts.append(ranges[curColumn + str(processingRowIndexNumber)].value)

            imgPath = "{}/figPyYcfg{}.png".format(PATH_TO_PYPLOT_IMAGES, diagramCounter)
            createBarAndSave(groups, counts, imgPath)

            gotData[i][1] = imgPath
        elif gotData[i][0] == "[двойнаястолбчатаядиаграмма]":
            gotData[i][0] = "Выплаты в прамках комплексного проекта"
            diagramCounter += 1

            # Damn.
            namesOfColumnsInDiagram = gotData[i][1].split("-")
            endLetter = namesOfColumnsInDiagram[1]
            namesOfColumnsInDiagram = [namesOfColumnsInDiagram[0]]
            curLetter = namesOfColumnsInDiagram[0]
            # print(endLetter)
            curLetter = column_index_from_string(curLetter)
            # print(curLetter, column_index_from_string(endLetter))
            while curLetter < column_index_from_string(endLetter):
                curLetter += 1
                namesOfColumnsInDiagram.append(get_column_letter(curLetter))

            groups = []
            fcounts = []
            scounts = []
            curColumn = 0
            while curColumn < len(namesOfColumnsInDiagram):
                fcounts.append(ranges[namesOfColumnsInDiagram[curColumn] + str(processingRowIndexNumber)].value)
                curColumn += 1
                groups.append(ranges[namesOfColumnsInDiagram[curColumn] + str(headersStartIndexNumber)].value)
                scounts.append(ranges[namesOfColumnsInDiagram[curColumn] + str(processingRowIndexNumber)].value)
                curColumn += 1

            imgPath = "{}/figPyYcfg{}.png".format(PATH_TO_PYPLOT_IMAGES, diagramCounter)
            for k in range(len(groups)):
                groups[k] = groups[k][-9:-4]
            createDoubleBarAndSave(groups, fcounts, scounts, imgPath)

            gotData[i][1] = imgPath
        elif gotData[i][0] == "[карта]":
            gotData[i][0] = ranges[gotData[i][1] + str(headersStartIndexNumber)].value
            gotInfo = ranges[gotData[i][1] + str(processingRowIndexNumber)].value
            gotData[i][1] = "{}/{}.png".format(PATH_TO_SUBJECTS, gotInfo)
        i += 1

    # print(gotData)

    gotData.insert(0, ranges[rules["startOfTable"][0] + str(processingRowIndexNumber)].value)

    # print("After")
    # print(gotData)

    return gotData


'''
    Sets paragraph's properties to plain text
'''


def setParaFormatPlainText(format, font):
    format.first_line_indent = Cm(1.27)
    format.line_spacing = 1
    format.space_after = Pt(0)

    format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    font.name = "Times New Roman"
    font.size = Pt(14)


'''
    Sets paragraph's properties to header
'''


def setParaFormatHeading(format, font):
    setParaFormatPlainText(format, font)
    font.bold = True


'''
    Sets paragraph's properties to analytics
'''


def setParaFormatAnal(format, font):
    format.line_spacing = 1
    format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    font.underline = True
    font.name = "Times New Roman"
    font.bold = True
    font.size = Pt(16)


'''
    Sets paragraph's properties to title text
'''


def setParaFormatTitle(format, font):
    format.line_spacing = 1
    format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    format.keep_together = True
    font.name = "Times New Roman"
    font.bold = True
    font.size = Pt(16)


'''
    Adds a picture to the paragraph
'''


def setPictureFormat(format, run, picPath):
    format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run.add_picture(picPath, width = Cm(17))


'''
    Forms .docx report from given information
'''


def formDocxFile(gotData, savePath):
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(1)
    filename = "Отчёт по {}".format(gotData[0])
    for s in ['*', '.', '/', '\\', '[', ']', ':', ';', '|', ',']:
        filename = filename.replace(s, '')
    filename += ".docx"
    filename = filename.replace('"', "").replace("»", "").replace("«", "").replace("\n", "")
    filename = savePath + "/" + filename

    titleImage = doc.add_paragraph()
    # print(os.getcwd())
    add_float_picture(titleImage, PATH_TO_MPT_LOGO, Cm(5), None, Cm(15.34), Cm(1.49))

    companyPic = PATH_TO_COMPANIES_LOGOS + "/" + gotData[0].replace('"', "")
    if not (os.path.exists(companyPic + ".jpg") or os.path.exists(companyPic + ".png") or os.path.exists(
            companyPic + ".png")):
        add_float_picture(titleImage, PATH_TO_MPT_LOGO, Cm(5), None, Cm(2.26),
                          Cm(1.49))  # Добавляем только в том случае, если
    else:
        add_float_picture(titleImage, PATH_TO_COMPANIES_LOGOS + "/" + gotData[0].replace('"', "") + ".jpg", Cm(2.5),
                          None,
                          Cm(2.26), Cm(1.49))
    titlePara = doc.add_paragraph()
    setParaFormatAnal(titlePara.paragraph_format, titlePara.add_run("АНАЛИТИЧЕСКИЙ ОТЧЁТ").font)
    titlePara = doc.add_paragraph()
    setParaFormatTitle(titlePara.paragraph_format,
                       titlePara.add_run("по комплексному проекту {}".format(
                           gotData[0].replace('"', "«", 1).replace('"', "»", 1))).font)
    gotData.pop(0)
    # print(gotData)
    for currentDataPair in gotData:
        # Oh Satan...
        if PATH_TO_PYPLOT_IMAGES in str(currentDataPair[1]) or PATH_TO_SUBJECTS in str(currentDataPair[1]):
            headerPara = doc.add_paragraph()
            setParaFormatHeading(headerPara.paragraph_format, headerPara.add_run("{}:".format(currentDataPair[0])).font)

            picturePara = doc.add_paragraph()
            setPictureFormat(picturePara.paragraph_format, picturePara.add_run(), currentDataPair[1])
        elif currentDataPair[0][:4] == "Дата":
            headerPara = doc.add_paragraph()
            setParaFormatHeading(headerPara.paragraph_format,
                                 headerPara.add_run("{}: {}".format(currentDataPair[0], currentDataPair[1])).font)
        else:
            headerPara = doc.add_paragraph()
            setParaFormatHeading(headerPara.paragraph_format,
                                 headerPara.add_run("{}: ".format(currentDataPair[0])).font)

            setParaFormatPlainText(headerPara.paragraph_format,
                                   headerPara.add_run("{}".format(str(currentDataPair[1]).replace("\n", " "))).font)

    try:
        doc.save(filename)
    except Exception:
        print()
        return PERMISSION_ERROR
    return 0


'''
    REALLY FUCKING IMPORTANT!!!
    We have a small trouble using some names with spaces. That's why
    we MUST send ways in quotes like 'C:/Go Fuck Yourself/Program Files/fuck fuck fuck fuck'
    Why can't we use '_'? BECAUSE OF FUCKING 'FUCK:/GO_FUCK/FUCKING_FUCK/FUCK FUCK FUCK'
    
    I think you've got it, right?
'''


def main():
    gotArgs = sys.argv

    try:
        tablePath = gotArgs[1]  # path to table to get info from
        # print(tablePath)
        docSavePath = gotArgs[2]  # path to save document
        # print(docSavePath)
        configPath = gotArgs[3]  # path to configuration
        # print(configPath)
        workRowNumber = int(gotArgs[4])  # number of row from table
        # print(workRowNumber)
    except Exception:
        return INVALID_ARGUMENTS_ERROR

    dictOfRulesForDoc = getRulesFromConfig(configPath)
    gotInfo = getInfoFromExcelTableUsingRules(tablePath, dictOfRulesForDoc, workRowNumber)

    return formDocxFile(gotInfo, docSavePath)


sys.exit(main())
